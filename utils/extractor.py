import PyPDF2
import docx
import os
import io

def extract_text(file_path):
    """Safely extracts text from PDF or DOCX file."""
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    try:
        if ext == '.pdf':
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
        elif ext == '.docx':
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip()
    except Exception as e:
        return f"Error: {str(e)}"

def is_resume(text):
    """
    Validation heuristic to check if the text contents are likely from a resume.
    Checks for presence of at least 2 common resume section keywords.
    """
    if not text:
        return False
    resume_keywords = ["experience", "education", "skills", "projects", "summary", "employment", "achievements"]
    text_lower = text.lower()
    matches = sum(1 for kw in resume_keywords if kw in text_lower)
    return matches >= 2

def extract_text_from_bytes(uploaded_file, filename):
    ext = os.path.splitext(filename)[1].lower()
    text = ""
    try:
        if ext == '.pdf':
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages:
                text += page.extract_text() or ""
        elif ext == '.docx':
            doc = docx.Document(uploaded_file)
            text = "\n".join([para.text for para in doc.paragraphs])
        
        extracted_text = text.strip()
        
        # Validation
        if not is_resume(extracted_text):
            return {"error": "Invalid Profile: The uploaded file does not appear to be a resume. Please upload a valid PDF or DOCX resume."}
            
        return extracted_text
    except Exception as e:
        return {"error": f"Corrupted or Unreadable File: {str(e)}"}
